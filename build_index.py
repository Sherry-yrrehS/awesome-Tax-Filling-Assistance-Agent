# 独立索引构建（带日志、异常）
from langchain_text_splitters import RecursiveCharacterTextSplitter
from variable import VALUE_ADDED_TAX_POLICY_FILE_WORD,COLLECTION_NAME1
from model import Embedding_model
from database import client
import re
from langchain_core.documents import Document
from docx import Document as DocxDocument
from logger_config import setup_logger

logger = setup_logger("build_index", log_file="build_index.log")

#加载文档
def load_documents():
    """加载所有文档，若单个文件失败则跳过"""
    all_docs = []
    logger.info("开始加载文档...")
    for file_path in VALUE_ADDED_TAX_POLICY_FILE_WORD:
        logger.info(f"  正在加载 {file_path} ...")
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]  #自动剔除 Word 里大量空行、空白段落，避免生成很多无用换行。
            content = "\n".join(paragraphs)
            if not content.strip():
                logger.warning(f"  文件 {file_path} 内容为空，跳过")
                continue
            doc_lc = Document(page_content=content, metadata={"source": file_path})
            all_docs.append(doc_lc)
            logger.info(f"  加载完成，共 {len(paragraphs)} 个非空段落")
        except FileNotFoundError:
            logger.error(f"  文件不存在: {file_path}")
        except Exception as e:
            logger.error(f"  加载 {file_path} 失败: {e}", exc_info=True)
    logger.info(f"文档加载完毕，成功加载 {len(all_docs)} 个文档")
    return all_docs

#自定义分割函数：优先按法律结构分割，保留条款完整性。
def split_documents_by_structure(documents, max_chunk_size=1500, overlap=200):
    """
    优先按法律结构分割，保留条款完整性。
    支持：
      - “第X条”
      - “一、”、“（一）”等政策编号
      - 章节标题（“第一章”）
    """
    final_chunks = []
    # 定义多个分割模式（按优先级）
    # 模式1: 第X条（最优先）
    pattern_article = re.compile(r'(?=\n第[零一二三四五六七八九十百]+条)')
    # 模式2: 一、二、三、...（中文数字加顿号）
    pattern_heading = re.compile(r'(?=\n[一二三四五六七八九十]+、)')
    # 模式3: （一）（二）...（带括号）
    pattern_subheading = re.compile(r'(?=\n（[一二三四五六七八九十]+）)')

    for doc in documents:
        text = doc.page_content
        original_meta = doc.metadata.copy()

        # 第一步：按“第X条”切分
        parts = re.split(pattern_article, text) #re.split(正则对象, 文本)：按照正则匹配到的内容切割字符串，把长文本切分成列表，结果存入变量parts。
        # 如果parts长度>1，说明有“第X条”
        if len(parts) > 1:
            # 处理每个部分
            for part in parts:
                part = part.strip() #strip() 是字符串自带方法，作用：删除字符串首尾所有空白字符
                if not part:
                    continue
                # 判断是否以“第X条”开头
                if re.match(r'^第[零一二三四五六七八九十百]+条', part):#从字符串开头匹配以「第 XX 条」开头的法条文本，匹配成功返回匹配对象，匹配失败返回 None。
                    # 检查长度，超长时二次切分（但保持条款完整？如果超长，按段落切分）
                    if len(part) > max_chunk_size:
                        #split_by_paragraphs返回splitter.split_documents([Document(page_content=text, metadata=metadata)])
                        sub_chunks = split_by_paragraphs(part, max_chunk_size, overlap, original_meta)
                        final_chunks.extend(sub_chunks)#将处理好的chunk内容增加到总chunk列表中（final_chunks)
                    else:
                        # 直接加入，记录元数据
                        new_meta = original_meta.copy()
                        new_meta["chunk_type"] = "article"
                        final_chunks.append(Document(page_content=part, metadata=new_meta))
                else:
                    # 可能是章标题或引言，单独处理（归入“其他”）
                    # 如果过长，也切分
                    if len(part) > max_chunk_size:
                        sub_chunks = split_by_paragraphs(part, max_chunk_size, overlap, original_meta)
                        final_chunks.extend(sub_chunks)
                    else:
                        new_meta = original_meta.copy()
                        new_meta["chunk_type"] = "other"
                        final_chunks.append(Document(page_content=part, metadata=new_meta))
        else:
            # 没有“第X条”，尝试按“一、”切分
            parts = re.split(pattern_heading, text)
            if len(parts) > 1:
                for part in parts:
                    part = part.strip()
                    if not part:
                        continue
                    if len(part) > max_chunk_size:
                        sub_chunks = split_by_paragraphs(part, max_chunk_size, overlap, original_meta)
                        final_chunks.extend(sub_chunks)
                    else:
                        new_meta = original_meta.copy()
                        new_meta["chunk_type"] = "heading"
                        final_chunks.append(Document(page_content=part, metadata=new_meta))
            else:
                # 再尝试按“（一）”切分
                parts = re.split(pattern_subheading, text)
                if len(parts) > 1:
                    for part in parts:
                        part = part.strip()
                        if not part:
                            continue
                        if len(part) > max_chunk_size:
                            sub_chunks = split_by_paragraphs(part, max_chunk_size, overlap, original_meta)
                            final_chunks.extend(sub_chunks)
                        else:
                            new_meta = original_meta.copy()
                            new_meta["chunk_type"] = "subheading"
                            final_chunks.append(Document(page_content=part, metadata=new_meta))
                else:
                    # 实在没有明显结构，直接用递归切分器
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=max_chunk_size,
                        chunk_overlap=overlap,
                        separators=["\n\n", "\n", "。", "；", " ", ""],
                        keep_separator=True,
                        strip_whitespace=True
                    )
                    sub_chunks = splitter.split_documents([doc])
                    for sc in sub_chunks:
                        sc.metadata["chunk_type"] = "fallback"
                    final_chunks.extend(sub_chunks)

    return final_chunks

#按段落二次切分的辅助函数
def split_by_paragraphs(text, max_size, overlap, metadata):
    """辅助函数：按段落二次切分"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", "。", "；", " ", ""],
        keep_separator=True,
        strip_whitespace=True
    )
    temp_doc = Document(page_content=text, metadata=metadata)#page_content：固定参数名，代表文档正文内容；metadata：元数据，字典格式，用来存这条文本的附加信息，不参与向量计算，但检索时可筛选过滤。
    return splitter.split_documents([temp_doc])


def main():
    try:
        # 1. 加载
        all_documents = load_documents()
        if not all_documents:
            logger.error("没有加载到任何文档，退出")
            return

        # 2. 切分
        logger.info("开始切割文档...")
        chunks = split_documents_by_structure(all_documents, max_chunk_size=1500, overlap=200)
        logger.info(f"文档共切分为 {len(chunks)} 个 chunk")

        # 3. 向量化
        logger.info("开始生成向量...")
        texts = [chunk.page_content for chunk in chunks]
        try:
            # 若文本数量很大，建议分批，这里简单处理
            vectors = Embedding_model.embed_documents(texts)
            logger.info("向量生成完成")
        except Exception as e:
            logger.error(f"向量化失败: {e}", exc_info=True)
            return

        # 4. 构建数据
        data = [
            {
                "id": i,
                "vector": vectors[i],
                "text": chunks[i].page_content,
                "source": chunks[i].metadata.get("source", "unknown"),
                "chunk_id": i,
            }
            for i in range(len(chunks))
        ]

        # 5. 插入 Milvus（含集合重建逻辑）
        logger.info("开始插入数据到 Milvus...")
        # 注意：database.py 现在不自动删除，所以在此主动重建（如需全量覆盖）
        from variable import EMBED_DIM
        if client.has_collection(COLLECTION_NAME1):
            client.drop_collection(COLLECTION_NAME1)
            logger.info("已删除旧集合")
        client.create_collection(
            collection_name=COLLECTION_NAME1,
            dimension=EMBED_DIM,
            metric_type="COSINE"
        )
        logger.info("新集合创建完成")

        insert_res = client.upsert(collection_name=COLLECTION_NAME1, data=data)
        client.flush(collection_name=COLLECTION_NAME1)
        logger.info("数据插入完成")

        # 6. 统计信息
        stats = client.get_collection_stats(collection_name=COLLECTION_NAME1)
        logger.info(f"集合统计信息: {stats}")
        results = client.query(
            collection_name=COLLECTION_NAME1,
            filter="id >= 0",
            output_fields=["id", "chunk_id"]
        )
        logger.info(f"当前集合共有 {len(results)} 条记录")

    except Exception as e:
        logger.critical(f"构建索引过程出现未捕获异常: {e}", exc_info=True)

if __name__ == "__main__":
    main()